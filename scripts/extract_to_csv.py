"""
Volunteer Data Onboarding - Extraction Pipeline.

This script pulls volunteer onboarding Microsoft Word documents (.docx) from Google Drive
and converts the structured form data into two neat, searchable CSV databases.
It handles highlighting, grid-based cell joining, and strict label matching.
"""

import os
import io
import csv
import argparse
import logging
from docx import Document
from pathlib import Path
from typing import Dict, List, Optional, Set


# GOOGLE DRIVE API IMPORTS

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload,MediaFileUpload


DEFAULT_INPUT_DIR = Path(__file__).parent.parent / "templates"
DEFAULT_OUTPUT_DIR = Path(__file__).parent.parent / "output"
DEFAULT_OUTPUT_FILENAME = "volunteer_onboarding.csv"
DEFAULT_FULL_FILENAME = "volunteer_onboarding_full.csv"

# Google Drive API Scopes
SCOPES = ["https://www.googleapis.com/auth/drive"]

# FIELD_MAP: Defines the essential columns for the primary "Filtered" CSV output.
FIELD_MAP = {
    "dubs impact driver (did) no.": "DID Number",
    "country": "Country",
    "province / state": "Province/State",
    "street address": "Street Address",
    "suburb": "Suburb/Area",
}

# FULL_FIELD_MAP: The master definition for all fields captured in the "Full" CSV.
FULL_FIELD_MAP = {
    # 1. Personal Information
    "dubs impact driver (did) no.": "DID Number",
    "first name": "First Name",
    "last name": "Last Name",
    "second name": "Second Name",
    "gender": "Gender",
    "phone number": "Phone Number",
    "race": "Race",
    "citizenship": "Citizenship",
    "foreign citizenship": "Foreign Citizenship",
    "home language": "Home Language",
    "date of birth (yyyymmdd)": "Date of Birth",
    "pronouns (optional)": "Pronouns",
    "name pronunciation(optional)": "Name Pronunciation",

    # 2. Contact Information
    "primary cellphone number": "Primary Cellphone Number",
    "whatsapp cellphone number": "WhatsApp Cellphone Number",
    "alternate cellphone number": "Alternate Cellphone Number",
    "personal email address": "Personal Email Address",

    # 3. Location Information
    "country": "Country",
    "province / state": "Province/State",
    "city": "City",
    "street address": "Street Address",
    "suburb / area": "Suburb/Area",

    # 4. Next of Kin
    "next of kin full name": "Next of Kin Full Name",
    "next of kin cellphone number": "Next of Kin Cellphone Number",

    # 5. Education
    "highest qualification": "Highest Qualification",
    "institution": "Institution",
    "qualification name": "Qualification Name",
    "year obtained": "Year Obtained",
    "former secondary/high school name": "Former School Name",

    # 6. Certifications
    "certification 001 - certification name": "Certification 001 - Certification Name",
    "certification 001 - issuing institution / provider": "Certification 001 - Issuing Institution / Provider",
    "certification 001 - year obtained (optional)": "Certification 001 - Year Obtained (optional)",
    "certification 002 - certification name": "Certification 002 - Certification Name",
    "certification 002 - issuing institution / provider": "Certification 002 - Issuing Institution / Provider",
    "certification 002 - year obtained (optional)": "Certification 002 - Year Obtained (optional)",
    "certification 003 - certification name": "Certification 003 - Certification Name",
    "certification 003 - issuing institution / provider": "Certification 003 - Issuing Institution / Provider",
    "certification 003 - year obtained (optional)": "Certification 003 - Year Obtained (optional)",
    "certification 004 - certification name": "Certification 004 - Certification Name",
    "certification 004 - issuing institution / provider": "Certification 004 - Issuing Institution / Provider",
    "certification 004 - year obtained (optional)": "Certification 004 - Year Obtained (optional)",
    "certification 005 - certification name": "Certification 005 - Certification Name",
    "certification 005 - issuing institution / provider": "Certification 005 - Issuing Institution / Provider",
    "certification 005 - year obtained (optional)": "Certification 005 - Year Obtained (optional)",

    # 7. Work / Role
    "team": "Team",
    "division": "Division",
    "role": "Role",
    "reporting officer/line manager": "Reporting Manager",
    "date joined (yyyymmdd)": "Date Joined",

    # 8. Skills & Interests
    "top technical skills": "Top Technical Skills",
    "top soft skills": "Top Soft Skills",
    "favourite quote or line": "Favourite Quote",
    "hobbies & interests": "Hobbies & Interests",
}

# Logger setup
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S"
)
logger = logging.getLogger(__name__)



# GOOGLE DRIVE API FUNCTIONS


def authenticate_drive():
    """Logs into Google Drive and returns the service object."""
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)

        with open('token.json', 'w') as token:
            token.write(creds.to_json())

    return build('drive', 'v3', credentials=creds)


def download_docs_from_drive(service, folder_id: str, download_dir: Path):
    """Downloads all Word docs from a Drive folder to the local directory."""
    download_dir.mkdir(parents=True, exist_ok=True)
    query = f"'{folder_id}' in parents and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"

    results = service.files().list(q=query, fields="files(id,name)").execute()
    items = results.get('files', [])

    if not items:
        logger.warning("No word documents found in Google Drive!!...")
        return

    for item in items:
        file_id = item['id']
        file_name = item['name']
        local_path = download_dir / file_name

        logger.info(f"Downloading from drive .......: {file_name}")
        # MERGED BUGFIX: Changed .file() to .files()
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()

        downloader = MediaIoBaseDownload(fh, request)

        done = False
        while done is False:
            # MERGED BUGFIX: Changed .next_chuck() to .next_chunk()
            status, done = downloader.next_chunk()

        with open(local_path, 'wb') as f:
            f.write(fh.getvalue())



def upload_csv_to_drive(service, file_path: Path, folder_id: str):
    file_name = file_path.name
    logger.info(f"Uploading csv to google drive .......: {file_name}")
    query = f"'{folder_id}' in parents and name='{file_name}' and trashed=false"
    results = service.files().list(q=query, fields="files(id, name)").execute()
    items = results.get('files', [])

    #prepare the local flie to be sent to internet
    media = MediaFileUpload(str(file_path), mimetype='text/csv', resumable=True)
    if items:
            # 3a. If it exists, UPDATE the existing file with the new data
            file_id = items[0]['id']
            service.files().update(fileId=file_id, media_body=media).execute()
            logger.info(f"Successfully UPDATED {file_name} in Google Drive!")
    else:
            #  If it does not exist, CREATE a brand new file
            file_metadata = {'name': file_name, 'parents': [folder_id]}
            service.files().create(body=file_metadata, media_body=media).execute()
            logger.info(f"Successfully UPLOADED new {file_name} to Google Drive!")
 



def is_cell_highlighted(cell) -> bool:
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                return True
    return False

def _process_table_row(row, context: Dict[str, str]) -> List[str]:
    row_lines = []
    cells_info = []

    for cell in row.cells:
        text = cell.text.strip().replace('\n', ' ')
        if text and (not cells_info or cells_info[-1]['text'] != text):
            cells_info.append({
                'text': text,
                'is_highlighted': is_cell_highlighted(cell)
            })

    if not cells_info:
        return row_lines

    cells_text = [c['text'] for c in cells_info]
    highlighted_cells = [c['text'] for c in cells_info if c['is_highlighted']]

    if len(cells_info) > 1:
        label_lower = cells_info[0]['text'].lower()
        if any(key in label_lower for key in ["citizenship", "gender", "race"]):
            if highlighted_cells:
                row_lines.append(f"{cells_info[0]['text']}: {highlighted_cells[0]}")
                return row_lines

    if len(cells_text) >= 3:
        label_0 = cells_info[0]['text'].lower()
        if "date" in label_0 or "did" in label_0:
            value = "".join([c['text'] for c in cells_info[1:] if c['text'] != ' ']).strip()
            row_lines.append(f"{cells_info[0]['text']}: {value}")
            return row_lines

        context['label'] = cells_text[0]
        sub_label = cells_text[1]
        value = " ".join(cells_text[2:])
        if highlighted_cells and highlighted_cells[0] != context['label']:
            value = highlighted_cells[0]

        row_lines.append(f"{context['label']} - {sub_label}: {value}")

    elif len(cells_text) == 2:
        label, value = cells_text[0], cells_text[1]
        if context['label'] and label in ["Certification Name", "Issuing Institution / Provider"]:
            row_lines.append(f"{context['label']} - {label}: {value}")
        else:
            if label.startswith("Certification"):
                context['label'] = label
            row_lines.append(f"{label}: {value}")

    elif len(cells_text) == 1:
        label = cells_text[0]
        if label.startswith("Certification"):
            context['label'] = label
        row_lines.append(label)

    return row_lines


def extract_text_from_docx(docx_path: Path) -> List[str]:
    lines = []
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        context = {'label': ""}
        for table in doc.tables:
            for row in table.rows:
                lines.extend(_process_table_row(row, context))

    except Exception as e:
        logger.error(f"Failed to read {docx_path.name}: {e}")

    return lines



# DATA STRUCTURING & CSV LOGIC


def parse_volunteer_data(lines: List[str], filename: str) -> Dict[str, str]:
    data = {field: "null" for field in FULL_FIELD_MAP.values()}

    for line in lines:
        if ":" not in line:
            continue

        parts = line.split(":", 1)
        raw_label = parts[0].strip().lower()
        value = parts[1].strip()

        if not value:
            continue

        for label, field in FULL_FIELD_MAP.items():
            if raw_label.startswith(label.lower()) or raw_label.startswith(field.lower()):
                if data[field] == "null":
                    data[field] = value
                break

    return data


def save_to_csv(data_list: List[Dict[str, str]], output_path: Path, fieldnames: List[str], pad_columns: bool = True):
    if not data_list:
        return

    column_widths = {}
    if pad_columns:
        for field in fieldnames:
            max_w = len(str(field))
            for data in data_list:
                val_w = len(str(data.get(field, "null")))
                max_w = max(max_w, val_w)
            column_widths[field] = max_w

    try:
        with open(output_path, mode='w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames, extrasaction='ignore', quoting=csv.QUOTE_MINIMAL)

            header_row = {f: (f.ljust(column_widths[f]) if pad_columns else f) for f in fieldnames}
            writer.writerow(header_row)

            for data in data_list:
                row = {}
                for field in fieldnames:
                    val = str(data.get(field, "null"))
                    row[field] = val.ljust(column_widths[field]) if pad_columns else val
                writer.writerow(row)

        logger.info(f"Saved {len(data_list)} records to {output_path.name} (Neat mode: {pad_columns})")

    except (PermissionError, Exception) as e:
        logger.error(f"Error saving to {output_path.name}: {e}")


def load_existing_records(output_path: Path) -> List[Dict[str, str]]:
    records = []
    if not output_path.exists():
        return records

    try:
        with open(output_path, mode='r', encoding='utf-8') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                clean_row = {k.strip(): v.strip() for k, v in row.items()}
                records.append(clean_row)
    except Exception as e:
        logger.error(f"Error reading existing records: {e}")

    return records





def process_onboarding_directory(input_dir: Path, output_file: Path, full_output_file: Path):
    """Orchestrates the scanning, parsing, deduplication, and saving of all forms."""

    all_filtered_records = load_existing_records(output_file)
    all_full_records = load_existing_records(full_output_file)
    existing_dids = {r.get("DID Number", "").strip() for r in all_filtered_records if r.get("DID Number")}

    if existing_dids:
        logger.info(f"Loaded {len(existing_dids)} records for deduplication check.")

    docx_files = list(input_dir.glob("*.docx"))
    if not docx_files and not all_filtered_records:
        logger.warning(f"No .docx files found in {input_dir}.")
        return

    new_count = 0
    for docx_file in docx_files:
        lines = extract_text_from_docx(docx_file)
        data = parse_volunteer_data(lines, docx_file.name)

        did = data.get("DID Number", "null").strip()
        if did != "null" and did in existing_dids:
            continue

        if any(v != "null" for k, v in data.items() if k in FULL_FIELD_MAP.values()):
            all_filtered_records.append(data)
            all_full_records.append(data)
            if did != "null":
                existing_dids.add(did)
            new_count += 1
            logger.info(f"Extracted: {docx_file.name}")

    if all_filtered_records:
        save_to_csv(all_filtered_records, output_file, list(FIELD_MAP.values()))
        full_headers = list(dict.fromkeys(FULL_FIELD_MAP.values()))
        save_to_csv(all_full_records, full_output_file, full_headers)

        summary = f"Done! {new_count} new records added." if new_count > 0 else "CSV files refreshed."
        logger.info(summary)



def main():
    """CLI entry point for the extraction utility."""
    parser = argparse.ArgumentParser(description="Volunteer Data Extraction Utility")
    parser.add_argument("--input", default=str(DEFAULT_INPUT_DIR), help="Source directory for .docx forms")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT_DIR), help="Destination directory for CSVs")
    parser.add_argument("--filename", default=DEFAULT_OUTPUT_FILENAME, help="Filtered CSV filename")
    parser.add_argument("--full-filename", default=DEFAULT_FULL_FILENAME, help="Full dataset filename")

    args = parser.parse_args()
    input_dir = Path(args.input)
    output_dir = Path(args.output)

    # Ensure directories exist
    input_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)


    # GOOGLE API: FOLDER IDs

    # Where to download the blank Word docs from:
    DRIVE_INPUT_FOLDER_ID = '18AbrxOmnSpl8z_tDJe6DOfnswewq8wLV'

    # Where to upload the finished CSVs to:
    DRIVE_OUTPUT_FOLDER_ID = '1CdLigYxEv4JmUnKzSDfY8rT1T7imwY8r'

    logger.info("Authenticating with Google Drive...")
    service = authenticate_drive()


    # STEP 1: PULL FROM DRIVE

    logger.info("Fetching files from Google Drive...")
    download_docs_from_drive(service, DRIVE_INPUT_FOLDER_ID, input_dir)


    # STEP 2: PROCESS LOCALLY

    logger.info(f"Scanning downloaded files in {input_dir.absolute()}...")
    process_onboarding_directory(input_dir, output_dir / args.filename, output_dir / args.full_filename)


    # STEP 3: PUSH TO DRIVE

    logger.info("Pushing updated CSV databases to Google Drive...")

    filtered_csv_path = output_dir / args.filename
    full_csv_path = output_dir / args.full_filename

    # Only try to upload if the files were actually created/updated locally
    if filtered_csv_path.exists():
        upload_csv_to_drive(service, filtered_csv_path, DRIVE_OUTPUT_FOLDER_ID)

    if full_csv_path.exists():
        upload_csv_to_drive(service, full_csv_path, DRIVE_OUTPUT_FOLDER_ID)

    logger.info("Pipeline Complete!")

if __name__ == "__main__":
    main()

          

